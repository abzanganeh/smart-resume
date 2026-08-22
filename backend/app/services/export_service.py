from __future__ import annotations

import io
import re
from pathlib import Path

import structlog
from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.agent.phase3_postprocess import is_category_skill_line
from app.models.cover_letter import CoverLetterOutput
from app.models.session import Session
from app.models.userinfo import UserInfo
from app.services.contact_authority import authoritative_contact

log = structlog.get_logger()

_RESUME_PDF_CSS = """
    @page { size: Letter; margin: 0.6in 0.65in; }
    body { font-family: Georgia, serif; font-size: 10.5pt; color: #111; line-height: 1.45; }
    h1 { font-size: 18pt; margin: 0 0 2pt; }
    h2 { font-size: 11pt; border-bottom: 1px solid #555; padding-bottom: 2pt; margin: 12pt 0 4pt; }
    ul { margin: 2pt 0; padding-left: 14pt; }
    li { margin-bottom: 2pt; }
    p  { margin: 2pt 0; }
"""

_COVER_LETTER_PDF_CSS = """
    @page { size: Letter; margin: 0.75in; }
    body { font-family: Georgia, serif; font-size: 11pt; color: #111; line-height: 1.55; }
    p { margin-bottom: 12pt; text-align: justify; }
"""

_TEMPLATE_DIR = Path(__file__).parent.parent / "templates"
_jinja_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATE_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)

# Common placeholder names/emails that LLMs emit when they lack real data
# (kept for backwards-compat re-exports; canonical logic lives in contact_authority.py)


def _format_skills_for_export(skills: list[str]) -> list[str]:
    """Render skills as category lines when categorized, else comma-joined fallback."""
    if not skills:
        return []
    if any(is_category_skill_line(s) for s in skills):
        return [s.strip() for s in skills if s.strip()]
    return [", ".join(skills)]


def _visible_bullets(bullets: list[str]) -> list[str]:
    return [b for b in bullets if b and b.strip()]


def _authoritative_contact(
    llm_contact: object,
    user: UserInfo | None,
    *,
    account_email: str | None = None,
) -> dict:
    return authoritative_contact(
        llm_contact,
        user_info=user,
        account_email=account_email,
    )


def _resume_to_html(session: Session) -> str:
    template = _jinja_env.get_template("resume.html")
    output = session.phase3_output
    user = session.user_info
    contact = _authoritative_contact(output.contact, user)
    return template.render(
        contact=contact,
        summary=output.summary,
        skill_lines=_format_skills_for_export(output.skills),
        experience=output.experience,
        projects=output.projects,
        education=output.education,
        certifications=output.certifications,
        user=user,
    )


async def render_pdf(session: Session) -> bytes:
    """Render the tailored resume to PDF via WeasyPrint (pure Python, no browser needed)."""
    # Imported lazily: ``app.services.export`` pulls in the assembler, which
    # imports this module back.
    from app.services.export.weasyprint_safe import render_pdf_bytes

    return render_pdf_bytes(_resume_to_html(session), css=[_RESUME_PDF_CSS])


def render_docx(session: Session) -> bytes:
    """Render the tailored resume to DOCX via python-docx."""
    output = session.phase3_output
    user = session.user_info
    doc = Document()

    # Title / Contact
    contact = _authoritative_contact(output.contact, user)
    name = contact.get("name", "")
    email = contact.get("email", "")
    phone = contact.get("phone", "")
    linkedin = contact.get("linkedin", "")
    github = contact.get("github", "")
    location = contact.get("location", "")
    website = contact.get("website", "")

    doc.add_heading(name, level=0)
    if location:
        doc.add_paragraph(location)
    contact_line_parts = [p for p in [email, phone, linkedin, github, website] if p]
    doc.add_paragraph("  |  ".join(contact_line_parts))

    # Summary
    if output.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(output.summary)

    # Skills
    if output.skills:
        doc.add_heading("Skills", level=1)
        for line in _format_skills_for_export(output.skills):
            doc.add_paragraph(line)

    # Experience
    if output.experience:
        doc.add_heading("Experience", level=1)
        for entry in output.experience:
            p = doc.add_paragraph()
            p.add_run(f"{entry.title} — {entry.company}").bold = True
            p.add_run(f"  ({entry.dates})")
            for bullet in _visible_bullets(entry.bullets):
                doc.add_paragraph(bullet, style="List Bullet")

    # Projects
    if output.projects:
        doc.add_heading("Projects", level=1)
        for proj in output.projects:
            if isinstance(proj, dict):
                p = doc.add_paragraph()
                p.add_run(proj.get("name", "")).bold = True
                if proj.get("url"):
                    p.add_run(f"  {proj['url']}")
                for bullet in _visible_bullets([b for b in proj.get("bullets", []) if isinstance(b, str)]):
                    doc.add_paragraph(bullet, style="List Bullet")

    # Education
    if output.education:
        doc.add_heading("Education", level=1)
        for edu in output.education:
            year_str = f" ({edu.year})" if edu.year else ""
            doc.add_paragraph(f"{edu.degree} — {edu.institution}{year_str}")
            for bullet in _visible_bullets(edu.bullets):
                doc.add_paragraph(bullet, style="List Bullet")

    # Certifications
    if output.certifications:
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph(", ".join(output.certifications))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render_txt(session: Session, *, account_email: str | None = None) -> str:
    """Plain-text resume for copy-paste."""
    output = session.phase3_output
    user = session.user_info
    lines: list[str] = []

    contact = _authoritative_contact(output.contact, user, account_email=account_email)
    name = contact.get("name", "")
    email = contact.get("email", "")
    phone = contact.get("phone", "")
    linkedin = contact.get("linkedin", "")
    github = contact.get("github", "")
    location = contact.get("location", "")
    website = contact.get("website", "")

    contact_parts = [p for p in [email, phone, linkedin, github, website] if p]
    lines += [name.upper()]
    if location:
        lines.append(location)
    lines += [" | ".join(contact_parts), ""]

    if output.summary:
        lines += ["SUMMARY", "-------", output.summary, ""]

    if output.skills:
        lines += ["SKILLS", "------", *_format_skills_for_export(output.skills), ""]

    if output.experience:
        lines += ["EXPERIENCE", "----------"]
        for entry in output.experience:
            lines += [f"{entry.title} | {entry.company} | {entry.dates}"]
            for bullet in _visible_bullets(entry.bullets):
                lines.append(f"  • {bullet}")
            lines.append("")

    if output.projects:
        lines += ["PROJECTS", "--------"]
        for proj in output.projects:
            if isinstance(proj, dict):
                lines.append(proj.get("name", ""))
                for bullet in _visible_bullets([b for b in proj.get("bullets", []) if isinstance(b, str)]):
                    lines.append(f"  • {bullet}")
                lines.append("")

    if output.education:
        lines += ["EDUCATION", "---------"]
        for edu in output.education:
            year_str = f" ({edu.year})" if edu.year else ""
            lines.append(f"{edu.degree} — {edu.institution}{year_str}")
            for bullet in _visible_bullets(edu.bullets):
                lines.append(f"  • {bullet}")
        lines.append("")

    if output.certifications:
        lines += ["CERTIFICATIONS", "--------------", ", ".join(output.certifications), ""]

    return "\n".join(lines)


def _cover_letter_paragraphs(output: CoverLetterOutput) -> list[str]:
    text = output.body_plain.strip() or output.body_markdown.strip()
    return [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]


def _cover_letter_to_html(session: Session) -> str:
    template = _jinja_env.get_template("cover_letter.html")
    output = session.cover_letter_output
    if output is None:
        raise ValueError("No cover letter to export")
    user = session.user_info
    target_role = user.target_role if user else ""
    return template.render(
        paragraphs=_cover_letter_paragraphs(output),
        user=user,
        target_role=target_role,
    )


async def render_cover_letter_pdf(session: Session) -> bytes:
    """Render the cover letter to PDF via WeasyPrint."""
    from app.services.export.weasyprint_safe import render_pdf_bytes

    return render_pdf_bytes(
        _cover_letter_to_html(session), css=[_COVER_LETTER_PDF_CSS]
    )


def render_cover_letter_docx(session: Session) -> bytes:
    """Render the cover letter to DOCX via python-docx."""
    output = session.cover_letter_output
    if output is None:
        raise ValueError("No cover letter to export")
    user = session.user_info
    doc = Document()

    if user and user.name:
        title = doc.add_heading(user.name, level=0)
        title.runs[0].font.size = Pt(14)
    contact_bits = []
    if user:
        if user.email:
            contact_bits.append(user.email)
        if user.phone:
            contact_bits.append(user.phone)
    if contact_bits:
        doc.add_paragraph("  |  ".join(contact_bits))

    if user and user.target_role:
        p = doc.add_paragraph()
        p.add_run(f"Re: {user.target_role}").italic = True

    doc.add_paragraph("")

    for para in _cover_letter_paragraphs(output):
        doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render_cover_letter_txt(session: Session) -> str:
    """Plain-text cover letter for copy-paste."""
    output = session.cover_letter_output
    if output is None:
        raise ValueError("No cover letter to export")
    user = session.user_info
    lines: list[str] = []

    if user and user.name:
        lines.append(user.name)
    if user and user.email:
        lines.append(user.email)
    if user and user.phone:
        lines.append(user.phone)
    if user and user.target_role:
        lines.append(f"Re: {user.target_role}")
    if lines:
        lines.append("")

    lines.extend(_cover_letter_paragraphs(output))
    return "\n\n".join(lines)


def _slug(text: str) -> str:
    cleaned = re.sub(r"[^\w\-]+", "_", text.strip().lower())
    return (cleaned[:60] or "resume").strip("_")


def _candidate_name(session: Session) -> str:
    if session.phase3_output and session.phase3_output.contact:
        name = (session.phase3_output.contact.get("name") or "").strip()
        if name:
            return name
    if session.user_info and session.user_info.name:
        return session.user_info.name.strip()
    if session.resume_parsed and session.resume_parsed.contact.name:
        return session.resume_parsed.contact.name.strip()
    return ""


def export_attachment_filename(session: Session, ext: str) -> str:
    """Build a safe download filename: company when tailoring to a JD, else candidate name."""
    from app.services.dashboard.resume_record import resolve_company_name

    normalized_ext = ext.lstrip(".")
    has_jd = bool((session.jd_raw or "").strip())
    if has_jd:
        company = resolve_company_name(session)
        if company and company not in ("Unknown", "—"):
            return f"{_slug(company)}_resume.{normalized_ext}"

    name = _candidate_name(session)
    slug = _slug(name) if name else "resume"
    return f"{slug}_resume.{normalized_ext}"
